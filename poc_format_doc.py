from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import os

def extract_highlighted_text_with_images_in_order(docx_path, output_dir='output_images'):
    document = Document(docx_path)
    os.makedirs(output_dir, exist_ok=True)
    

    iter_block_items(document)

    

# 取得段落、表格中所有文字區段（含高亮）與圖片（二選一）
def iter_block_items(document):
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for index, run in enumerate(para.runs):
                        if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                            text = run.text.strip()
                            if (text):
                                print(text)
    """
    for rel in document.part._rels:
        rel_obj = document.part._rels[rel]
        if "image" in rel_obj.target_ref:
            yield rel_obj.target_part.blob
    """

# 使用範例
docx_path = os.path.join(os.getcwd(), "processed_document", '字庫A.docx')
pairs = extract_highlighted_text_with_images_in_order(docx_path)

for text, image in pairs:
    print(f"🟡 高亮文字: {text} -> 🖼️ 圖片: {image}")
