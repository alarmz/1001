from docx import Document
from docx.shared import RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

# 設定參數
target_chars = ['張', '員', '瑞']
image_path = 'avatar.png'  # 圖片路徑
input_path = '1001.docx'
output_path = '1001-1.docx'

# 開啟文件
doc = Document(input_path)

# 處理每一段
for para in doc.paragraphs:
    new_runs = []
    for run in para.runs:
        new_text = ""
        for char in run.text:
            if char in target_chars:
                # 建立新的 run 並設定 highlight
                new_run = para.add_run(char)
                new_run.font.highlight_color = 7  # 黃色 Highlight
                new_runs.append(new_run)

                # 加入圖片
                run_after_img = para.add_run()
                run_after_img.add_picture(image_path, width=Inches(0.2))
            else:
                new_run = para.add_run(char)
                new_runs.append(new_run)

        # 清空原來的文字
        run.text = ""

# 儲存新文件
doc.save(output_path)
print(f"處理完成，輸出至 {output_path}")


import sqlite3
from io import BytesIO
from docx import Document
from docx.shared import Inches

# 假設你的資料表結構是 (id INTEGER, image BLOB)
cursor.execute("SELECT image FROM images WHERE id = ?", (1,))
result = cursor.fetchone()

if result and result[0]:
    image_blob = result[0]  # 這是 BLOB 的 bytes

    # 包裝成 BytesIO 給 add_picture 用
    image_stream = BytesIO(image_blob)
    

    # 建立 Word 文件並加入圖片
    doc = Document()
    p = doc.add_paragraph("圖片來自 SQLite:")
    run = p.add_run()
    run.add_picture(image_stream, width=Inches(1.5))
    doc.save("from_sqlite.docx")
    print("圖片成功從 SQLite 加入 Word 文件！")
else:
    print("找不到圖片資料")

conn.close()    