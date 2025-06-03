import os
import uuid
import tempfile
import secrets
from pathlib import Path
from typing import Dict, List
import re
from docx import Document

from Create_Docx import  CreateDocx

TEMP_DIR = Path(tempfile.gettempdir()) / "nicegui_docs"
TEMP_DIR.mkdir(exist_ok=True)

POLYPHONE_DICT = {
    '行': ['xíng', 'háng'], '長': ['cháng', 'zhǎng'], '重': ['zhòng', 'chóng'],
    '數': ['shù', 'shǔ'], '分': ['fēn', 'fèn'], '中': ['zhōng', 'zhòng'],
    '為': ['wéi', 'wèi'], '間': ['jiān', 'jiàn'], '乾': ['gān', 'qián'],
    '血': ['xuè', 'xiě'], '角': ['jiǎo', 'jué'], '便': ['biàn', 'pián'],
    '調': ['tiáo', 'diào'], '量': ['liáng', 'liàng'], '教': ['jiāo', 'jiào'],
    '背': ['bèi', 'bēi'], '種': ['zhǒng', 'zhòng'], '少': ['shǎo', 'shào'],
    '還': ['hái', 'huán'], '了': ['le', 'liǎo']
}

VARIANT_DICT = {
    '台': '臺', '裏': '裡', '么': '麼', '着': '著', '説': '說', '綫': '線',
    '衆': '眾', '糰': '團', '麵': '面', '鞦韆': '秋千', '脩': '修', '衹': '只',
    '巖': '岩', '峯': '峰', '島': '島', '麽': '麼', '綵': '彩', '劃': '畫',
    '製': '制', '適': '適'
}

class DocumentProcessor:
    @staticmethod
    def text_to_docx(text_content: str, filename: str) -> str:
        #doc = Document()
        docx = CreateDocx()
        
        filepath = TEMP_DIR / filename
        docx.text_files.append(filepath)
        sDocxTextPath = docx.Save_Text_to_Docx_web(text_content, TEMP_DIR, filename)
        docx.Process_Docx_Word_by_word(sDocxTextPath)
        return str(sDocxTextPath)

    @staticmethod
    def find_polyphones(text: str) -> List[Dict]:
        polyphones = []
        for char, pronunciations in POLYPHONE_DICT.items():
            positions = []
            for i, c in enumerate(text):
                if c == char:
                    start = max(0, i-5)
                    end = min(len(text), i+6)
                    context = text[start:end]
                    positions.append({
                        'position': i,
                        'context': context,
                        'char_index': i - start
                    })
            if positions:
                polyphones.append({
                    'char': char,
                    'pronunciations': pronunciations,
                    'positions': positions
                })
        return polyphones

    @staticmethod
    def find_variants(text: str) -> List[Dict]:
        variants = []
        for variant, standard in VARIANT_DICT.items():
            positions = []
            for pos in re.finditer(re.escape(variant), text):
                start_pos = pos.start()
                start = max(0, start_pos-5)
                end = min(len(text), start_pos+len(variant)+5)
                context = text[start:end]
                positions.append({
                    'position': start_pos,
                    'context': context,
                    'char_index': start_pos - start
                })
            if positions:
                variants.append({
                    'variant': variant,
                    'standard': standard,
                    'positions': positions
                })
        return variants

    @staticmethod
    def process_docx_file(filepath: str) -> tuple:
        doc = Document(filepath)
        full_text = [para.text for para in doc.paragraphs]
        text_content = '\n'.join(full_text)

        polyphones = DocumentProcessor.find_polyphones(text_content)
        polyphone_doc = Document()
        polyphone_doc.add_heading('破音字檢查報告', 0)
        if polyphones:
            for item in polyphones:
                polyphone_doc.add_heading(f'破音字: {item["char"]}', level=1)
                polyphone_doc.add_paragraph(f'可能讀音: {", ".join(item["pronunciations"])}')
                polyphone_doc.add_paragraph('出現位置:')
                for pos in item['positions']:
                    p = polyphone_doc.add_paragraph()
                    p.add_run(f'位置 {pos["position"]}: ').bold = True
                    context = pos['context']
                    idx = pos['char_index']
                    p.add_run(context[:idx])
                    highlight = p.add_run(context[idx])
                    highlight.bold = True
                    p.add_run(context[idx+1:])
        else:
            polyphone_doc.add_paragraph('未發現破音字。')

        variants = DocumentProcessor.find_variants(text_content)
        variant_doc = Document()
        variant_doc.add_heading('異體字檢查報告', 0)
        if variants:
            for item in variants:
                variant_doc.add_heading(f'異體字: {item["variant"]} → {item["standard"]}', level=1)
                variant_doc.add_paragraph('出現位置:')
                for pos in item['positions']:
                    p = variant_doc.add_paragraph()
                    p.add_run(f'位置 {pos["position"]}: ').bold = True
                    context = pos['context']
                    idx = pos['char_index']
                    p.add_run(context[:idx])
                    highlight = p.add_run(item['variant'])
                    highlight.bold = True
                    p.add_run(context[idx+len(item['variant']):])
                    p.add_run(f' → 建議改為: {item["standard"]}').italic = True
        else:
            variant_doc.add_paragraph('未發現異體字。')

        polyphone_path = TEMP_DIR / f"todo_破音字_{uuid.uuid4().hex[:8]}.docx"
        variant_path = TEMP_DIR / f"todo_異體字_{uuid.uuid4().hex[:8]}.docx"
        polyphone_doc.save(polyphone_path)
        variant_doc.save(variant_path)

        return str(polyphone_path), str(variant_path)

class UserSession:
    def __init__(self):
        self.user_id = str(uuid.uuid4())
        self.uploaded_files = []
        self.generated_files = []

user_sessions: Dict[str, UserSession] = {}

def get_user_session():
    from nicegui import app
    if not hasattr(app.storage.user, 'session_id'):
        app.storage.user.session_id = str(uuid.uuid4())
        user_sessions[app.storage.user.session_id] = UserSession()
    return user_sessions[app.storage.user.session_id]

def cleanup_old_files():
    import time
    try:
        for file_path in TEMP_DIR.glob("*"):
            if file_path.is_file() and file_path.stat().st_mtime < (time.time() - 3600):
                file_path.unlink()
    except Exception as e:
        print(f"清理文件時出錯: {e}")
