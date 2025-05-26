import os
import uuid
import tempfile
import secrets
from pathlib import Path
from typing import Dict, List
import re

from nicegui import ui, app
from docx import Document
from docx.shared import Inches

# 創建臨時文件目錄
TEMP_DIR = Path(tempfile.gettempdir()) / "nicegui_docs"
TEMP_DIR.mkdir(exist_ok=True)

# 破音字數據庫（示例）
POLYPHONE_DICT = {
    '行': ['xíng', 'háng'],
    '長': ['cháng', 'zhǎng'],
    '重': ['zhòng', 'chóng'],
    '數': ['shù', 'shǔ'],
    '分': ['fēn', 'fèn'],
    '中': ['zhōng', 'zhòng'],
    '為': ['wéi', 'wèi'],
    '間': ['jiān', 'jiàn'],
    '乾': ['gān', 'qián'],
    '血': ['xuè', 'xiě'],
    '角': ['jiǎo', 'jué'],
    '便': ['biàn', 'pián'],
    '調': ['tiáo', 'diào'],
    '量': ['liáng', 'liàng'],
    '教': ['jiāo', 'jiào'],
    '背': ['bèi', 'bēi'],
    '種': ['zhǒng', 'zhòng'],
    '少': ['shǎo', 'shào'],
    '還': ['hái', 'huán'],
    '了': ['le', 'liǎo']
}

# 異體字對照表（示例）
VARIANT_DICT = {
    '台': '臺',
    '裏': '裡',
    '么': '麼',
    '着': '著',
    '説': '說',
    '綫': '線',
    '衆': '眾',
    '糰': '團',
    '麵': '面',
    '鞦韆': '秋千',
    '脩': '修',
    '衹': '只',
    '巖': '岩',
    '峯': '峰',
    '島': '島',
    '麽': '麼',
    '綵': '彩',
    '劃': '畫',
    '製': '制',
    '適': '適'
}

class DocumentProcessor:
    """文檔處理器"""
    
    @staticmethod
    def text_to_docx(text_content: str, filename: str) -> str:
        """將文字轉換為 docx 文件"""
        doc = Document()
        
        # 添加標題
        doc.add_heading('文字轉換文檔', 0)
        
        # 分段處理文本
        paragraphs = text_content.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para)
        
        # 保存文件
        filepath = TEMP_DIR / filename
        doc.save(filepath)
        return str(filepath)
    
    @staticmethod
    def find_polyphones(text: str) -> List[Dict]:
        """找出文本中的破音字"""
        polyphones = []
        for char, pronunciations in POLYPHONE_DICT.items():
            positions = []
            for i, c in enumerate(text):
                if c == char:
                    # 獲取上下文
                    start = max(0, i-5)
                    end = min(len(text), i+6)
                    context = text[start:end]
                    positions.append({
                        'position': i,
                        'context': context,
                        'char_index': i - start
                    })
            
            if positions:
                polyphones.append({
                    'char': char,
                    'pronunciations': pronunciations,
                    'positions': positions
                })
        
        return polyphones
    
    @staticmethod
    def find_variants(text: str) -> List[Dict]:
        """找出文本中的異體字"""
        variants = []
        for variant, standard in VARIANT_DICT.items():
            positions = []
            for i, pos in enumerate(re.finditer(re.escape(variant), text)):
                start_pos = pos.start()
                # 獲取上下文
                start = max(0, start_pos-5)
                end = min(len(text), start_pos+len(variant)+5)
                context = text[start:end]
                positions.append({
                    'position': start_pos,
                    'context': context,
                    'char_index': start_pos - start
                })
            
            if positions:
                variants.append({
                    'variant': variant,
                    'standard': standard,
                    'positions': positions
                })
        
        return variants
    
    @staticmethod
    def process_docx_file(filepath: str) -> tuple:
        """處理 docx 文件，生成破音字和異體字報告"""
        doc = Document(filepath)
        full_text = []
        
        # 提取所有文字
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        text_content = '\n'.join(full_text)
        
        # 處理破音字
        polyphones = DocumentProcessor.find_polyphones(text_content)
        polyphone_doc = Document()
        polyphone_doc.add_heading('破音字檢查報告', 0)
        
        if polyphones:
            for item in polyphones:
                polyphone_doc.add_heading(f'破音字: {item["char"]}', level=1)
                polyphone_doc.add_paragraph(f'可能讀音: {", ".join(item["pronunciations"])}')
                polyphone_doc.add_paragraph('出現位置:')
                
                for pos in item['positions']:
                    p = polyphone_doc.add_paragraph()
                    p.add_run(f'位置 {pos["position"]}: ').bold = True
                    
                    # 高亮顯示目標字符
                    context = pos['context']
                    char_idx = pos['char_index']
                    p.add_run(context[:char_idx])
                    highlight_run = p.add_run(context[char_idx])
                    highlight_run.bold = True
                    p.add_run(context[char_idx+1:])
                
                polyphone_doc.add_paragraph('')
        else:
            polyphone_doc.add_paragraph('未發現破音字。')
        
        # 處理異體字
        variants = DocumentProcessor.find_variants(text_content)
        variant_doc = Document()
        variant_doc.add_heading('異體字檢查報告', 0)
        
        if variants:
            for item in variants:
                variant_doc.add_heading(f'異體字: {item["variant"]} → {item["standard"]}', level=1)
                variant_doc.add_paragraph('出現位置:')
                
                for pos in item['positions']:
                    p = variant_doc.add_paragraph()
                    p.add_run(f'位置 {pos["position"]}: ').bold = True
                    
                    # 高亮顯示異體字
                    context = pos['context']
                    char_idx = pos['char_index']
                    p.add_run(context[:char_idx])
                    highlight_run = p.add_run(item['variant'])
                    highlight_run.bold = True
                    p.add_run(context[char_idx+len(item['variant']):])
                    
                    # 建議修改
                    p.add_run(f' → 建議改為: {item["standard"]}').italic = True
                
                variant_doc.add_paragraph('')
        else:
            variant_doc.add_paragraph('未發現異體字。')
        
        # 保存報告文件
        polyphone_path = TEMP_DIR / f"todo_破音字_{uuid.uuid4().hex[:8]}.docx"
        variant_path = TEMP_DIR / f"todo_異體字_{uuid.uuid4().hex[:8]}.docx"
        
        polyphone_doc.save(polyphone_path)
        variant_doc.save(variant_path)
        
        return str(polyphone_path), str(variant_path)

class UserSession:
    """用戶會話管理"""
    def __init__(self):
        self.user_id = str(uuid.uuid4())
        self.uploaded_files = []
        self.generated_files = []

# 全局用戶會話存儲
user_sessions: Dict[str, UserSession] = {}

def get_user_session() -> UserSession:
    """獲取當前用戶會話"""
    if not hasattr(app.storage.user, 'session_id'):
        app.storage.user.session_id = str(uuid.uuid4())
        user_sessions[app.storage.user.session_id] = UserSession()
    
    return user_sessions[app.storage.user.session_id]

def cleanup_old_files():
    """清理舊文件"""
    try:
        for file_path in TEMP_DIR.glob("*"):
            if file_path.is_file():
                # 刪除超過1小時的文件
                if file_path.stat().st_mtime < (time.time() - 3600):
                    file_path.unlink()
    except Exception as e:
        print(f"清理文件時出錯: {e}")

@ui.page('/')
def main():
    """主頁面"""
    session = get_user_session()
    
    ui.html('<h1 style="text-align: center; color: #2563eb;">📄 文檔處理系統</h1>')
    ui.html('<p style="text-align: center; color: #64748b;">支持文字轉 docx 和文檔內容分析</p>')
    
    with ui.tabs().classes('w-full') as tabs:
        tab1 = ui.tab('文字轉 docx', icon='text_fields')
        tab2 = ui.tab('docx 分析', icon='analytics')
    
    with ui.tab_panels(tabs, value=tab1).classes('w-full'):
        # 第一個功能：文字轉 docx
        with ui.tab_panel(tab1):
            ui.html('<h2>📝 文字轉 docx</h2>')
            ui.html('<p>上傳文字檔案或直接輸入文字，系統將生成 docx 文檔供您下載。</p>')
            
            with ui.card().classes('w-full max-w-4xl mx-auto'):
                # 文字輸入區域
                text_input = ui.textarea(
                    label='輸入文字內容',
                    placeholder='請輸入要轉換的文字內容...',
                    value=''
                ).classes('w-full').style('min-height: 200px;')
                
                # 文件上傳
                ui.separator()
                ui.label('或上傳文字檔案 (.txt, .md):')
                
                def handle_text_upload(e):
                    try:
                        # 檢查文件類型
                        if not e.name.lower().endswith(('.txt', '.md')):
                            ui.notify('請上傳 .txt 或 .md 文件！', type='warning')
                            return
                            
                        content = e.content.read().decode('utf-8')
                        text_input.value = content
                        ui.notify(f'文件 "{e.name}" 上傳成功！', type='positive')
                    except Exception as ex:
                        ui.notify(f'文件上傳失敗: {str(ex)}', type='negative')
                
                ui.upload(
                    on_upload=handle_text_upload,
                    auto_upload=True
                ).classes('w-full')
                
                ui.separator()
                
                # 轉換按鈕
                def convert_to_docx():
                    if not text_input.value.strip():
                        ui.notify('請輸入文字內容或上傳文件！', type='warning')
                        return
                    
                    try:
                        filename = f"converted_{uuid.uuid4().hex[:8]}.docx"
                        filepath = DocumentProcessor.text_to_docx(text_input.value, filename)
                        session.generated_files.append(filepath)
                        
                        # 提供下載
                        ui.download(filepath, filename)
                        ui.notify('文檔生成成功！', type='positive')
                        
                    except Exception as ex:
                        ui.notify(f'轉換失敗: {str(ex)}', type='negative')
                
                ui.button('🔄 轉換為 docx', on_click=convert_to_docx).classes('w-full')
        
        # 第二個功能：docx 分析
        with ui.tab_panel(tab2):
            ui.html('<h2>🔍 docx 文檔分析</h2>')
            ui.html('<p>上傳 docx 文檔，系統將分析其中的破音字和異體字，並生成詳細報告。</p>')
            
            with ui.card().classes('w-full max-w-4xl mx-auto'):
                upload_result = ui.label('').classes('mt-4')
                download_area = ui.column().classes('w-full mt-4')
                
                def handle_docx_upload(e):
                    try:
                        # 檢查文件類型
                        if not e.name.lower().endswith('.docx'):
                            ui.notify('請上傳 .docx 文件！', type='warning')
                            return
                            
                        # 保存上傳的文件
                        temp_path = TEMP_DIR / f"upload_{uuid.uuid4().hex[:8]}_{e.name}"
                        with open(temp_path, 'wb') as f:
                            f.write(e.content.read())
                        
                        upload_result.text = f'正在分析文件: {e.name}...'
                        upload_result.classes('text-blue-600')
                        
                        # 處理文檔
                        polyphone_path, variant_path = DocumentProcessor.process_docx_file(str(temp_path))
                        
                        # 清空下載區域
                        download_area.clear()
                        
                        with download_area:
                            ui.html('<h3>📥 分析結果下載</h3>')
                            with ui.row().classes('w-full gap-4'):
                                with ui.card().classes('flex-1'):
                                    ui.html('<h4>🔤 破音字報告</h4>')
                                    ui.html('<p>檢查文檔中可能存在讀音歧義的字詞</p>')
                                    ui.button(
                                        '下載破音字報告',
                                        on_click=lambda: ui.download(
                                            polyphone_path,
                                            Path(polyphone_path).name
                                        ),
                                        icon='download'
                                    ).classes('w-full')
                                
                                with ui.card().classes('flex-1'):
                                    ui.html('<h4>📝 異體字報告</h4>')
                                    ui.html('<p>檢查文檔中的異體字使用情況</p>')
                                    ui.button(
                                        '下載異體字報告',
                                        on_click=lambda: ui.download(
                                            variant_path,
                                            Path(variant_path).name
                                        ),
                                        icon='download'
                                    ).classes('w-full')
                        
                        upload_result.text = f'✅ 文件 "{e.name}" 分析完成！請下載報告查看結果。'
                        upload_result.classes('text-green-600')
                        
                        # 清理臨時上傳文件
                        os.unlink(temp_path)
                        
                    except Exception as ex:
                        upload_result.text = f'❌ 處理失敗: {str(ex)}'
                        upload_result.classes('text-red-600')
                
                ui.label('選擇 docx 文件上傳:')
                ui.upload(
                    on_upload=handle_docx_upload,
                    auto_upload=True,
                    max_file_size=10_000_000
                ).classes('w-full')
    
    # 頁腳信息
    ui.separator().classes('mt-8')
    ui.html('''
        <div style="text-align: center; padding: 20px; color: #64748b;">
            <p>🔧 多用戶文檔處理系統 | 支持並發使用</p>
            <p><small>支持的格式: .txt, .md, .docx | 文件大小限制: 10MB</small></p>
        </div>
    ''')

if __name__ in {"__main__", "__mp_main__"}:
    import time
    
    # 定期清理舊文件
    ui.timer(3600, cleanup_old_files)  # 每小時清理一次
    
    # 生成或獲取存儲密鑰
    storage_secret = os.getenv('NICEGUI_STORAGE_SECRET', secrets.token_urlsafe(32))
    
    ui.run(
        title='文檔處理系統',
        favicon='📄',
        port=8080,
        host='0.0.0.0',
        reload=False,
        show=True,
        storage_secret=storage_secret
    )