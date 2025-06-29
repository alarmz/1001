import os
from docx import Document
import sqlite3
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Inches
import tempfile
from io import BytesIO

from io import BytesIO
from PIL import Image, UnidentifiedImageError

class CreateDocx:
    def __init__(self):
        self.source_text = os.path.join(os.getcwd(), "source_text")
        self.conn = sqlite3.connect("word_data.db")
        self.conn.row_factory = sqlite3.Row
        self.cursor = self.conn.cursor()
        self.text_files = []
        pass
    
    def Search_in_Database(self, src_word):
        self.cursor.execute(f"SELECT * FROM Word where sWord = '{src_word}'")
        db_res = [dict(row) for row in self.cursor.fetchall()]
        return db_res
        
    
    def list_text_files(self):
        self.text_files = [f for f in os.listdir(self.source_text) if f.endswith('.txt')]
        
    def Save_Text_to_Docx(self, sPath, sFileName):
        sTextPath = os.path.join(self.source_text, sFileName)
        with open(sTextPath, "r", encoding="utf-8") as file:
            text = file.readlines()
            
        doc = Document()
        for line in text:
            doc.add_paragraph(line.strip())
            
        sFileName = sFileName.replace("txt", "docx")
        stage1_processed =  os.path.join(os.getcwd(), "stage1_processed")
        if (os.path.exists(stage1_processed) == False):
            os.makedirs(stage1_processed)        
        docxFileName = os.path.join(stage1_processed, sFileName)
        sDocxTextPath = os.path.join(sPath, docxFileName)
        doc.save(sDocxTextPath)
        return sDocxTextPath
    
    def Save_BIBO_to_temp_File(self, imgData):
        temp_img_path = ""
        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as temp_img:
            temp_img.write(imgData)
            temp_img_path = temp_img.name
        return temp_img_path
    
    def Mapping_DB_and_Source(self, src_a_Word):
        db_res = self.Search_in_Database(src_a_Word)
        return db_res
    
    def docx_add_yellow_highlight(self, run, db_res):
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        #img_path = self.Save_BIBO_to_temp_File(db_res[0]["imgData"])
        image_stream = BytesIO(db_res[0]["imgData"])
        run.add_picture(image_stream, width=Inches(1.0))
        run.text = ""
            
    def docx_add_Green_highlight(self, para):
        for run in para.runs:
            run.font.highlight_color =  WD_COLOR_INDEX.BRIGHT_GREEN
    
    def Process_Docx_Word_by_word(self, sDocxTextPath):
        doc = Document(sDocxTextPath)
        for para in doc.paragraphs:
            new_runs = []
            for run in para.runs:
                for ch in run.text:
                    # 在這裡可以加入比對或處理邏輯
                    #print(f"seek {ch}")
                    db_res = self.Mapping_DB_and_Source(ch.strip())
                    #print(ch)  # 每次印出一個字
                    if (db_res == []):
                        #print("This is new word, not in Database.")
                        new_run = para.add_run(ch)
                        #new_run.font.highlight_color = WD_COLOR_INDEX.GRAY_25  
                        new_runs.append(new_run)                        
                    else:
                        if (db_res[0]["sType"] == "normal"):#
                            new_run = para.add_run(ch)
                            new_runs.append(new_run)
                        #elif (db_res[0]["sType"] == "FontType") and (db_res[0]["isIgnore"] == 0):
                        elif (db_res[0]["sType"] in ["A1", "A2"]) and (db_res[0]["isIgnore"] == 0):
                            new_run = para.add_run(ch)
                            new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW  # 黃色 Highlight
                            new_runs.append(new_run)
                            # 加入圖片
                            run_after_img = para.add_run()
                            image_stream = BytesIO(db_res[0]["imgData"])
                            #self.Save_BIBO_to_temp_File(image_stream)
                            self.Save_BIBO_to_temp_File(image_stream.getvalue())  # ✅ 傳入 bytes
                            run_after_img.add_picture(image_stream, width=Inches(0.2))
                            print(f"{ch}...")
                        elif (db_res[0]["sType"] == "B") and (db_res[0]["isIgnore"] == 1):
                            new_run = para.add_run(ch)
                            new_run.font.highlight_color = WD_COLOR_INDEX.GRAY_25  # GRAY_25
                            new_runs.append(new_run)
                        elif (db_res[0]["sType"] in ["A3", "A4", "A5"]):
                            #a3差異字＋難字
                            new_run = para.add_run(ch)
                            new_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN  # BRIGHT_GREEN
                            new_runs.append(new_run)
                            # 加入圖片
                            run_after_img = para.add_run()
                            
                            #---------
                            raw_blob = db_res[0]["imgData"]
                            
                            try:
                                image = Image.open(BytesIO(raw_blob))
                                image.verify()  # 先驗證合法性
                                image = Image.open(BytesIO(raw_blob))  # 再次開啟（verify 會破壞 stream）
                                image_stream = BytesIO()
                                image.save(image_stream, format='PNG')
                                image_stream.seek(0)
                                run_after_img.add_picture(image_stream, width=Inches(0.2))
                            except UnidentifiedImageError:
                                pass
                                #print(f"❌ [{ch}]資料庫中的圖片不是合法格式")
                            except Exception as e:
                                print(f"❌ 加入圖片時出錯: {e}")                            
                            #-------
                            
                            
                            #image_stream = BytesIO(db_res[0]["imgData"])
                            #run_after_img.add_picture(image_stream, width=Inches(0.2))
                            print(f"{ch}...")
                        elif (db_res[0]["sType"] == "dual"):
                            new_run = para.add_run(ch)
                            new_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                            new_runs.append(new_run)
                            
                    
                run.text = ""        
        doc.save(sDocxTextPath)
    
        
        
    def Loop_Text_Files_Create_Docx(self):
        for aTxt in self.text_files:
            sFileFull_Path = os.path.join(os.getcwd(), "source_text")
            if (os.path.exists(sFileFull_Path) == False):
                os.makedirs(sFileFull_Path)            
            sDocxTextPath = self.Save_Text_to_Docx(sFileFull_Path, aTxt)
            
            #print(f"process=> {sFileFull_Path}")
            #print(f"Convert Doctx => {sDocxTextPath}----Done")
            #Open Docx
            self.Process_Docx_Word_by_word(sDocxTextPath)
            #Loop every words to mapping database 
            return sDocxTextPath

    
    
    



def main():
    docx =  CreateDocx()
    docx.list_text_files()
    docx.Loop_Text_Files_Create_Docx()


if __name__ == "__main__":
    main()