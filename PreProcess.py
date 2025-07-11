from docx import Document
from docx.oxml.ns import qn
import os
import sqlite3
import random
import string
from docx.enum.text import WD_COLOR_INDEX


class docx1001:
    def __init__(self):
        self.processed_data = os.path.join(".", "processed_document")
        self.docx_files = []
        self.conn = sqlite3.connect("word_data.db")
        self.cursor = self.conn.cursor()
        self.Create_Sqlite_DB()
        self.All_Processed_words = []
        self.YELLOW = 7
        self.BRIGHT_GREEN = 4
        self.pool = string.ascii_letters + string.digits
        pass
    
    def Create_Sqlite_DB(self):
        create_table_sql = """
        CREATE TABLE IF NOT EXISTS Word (
            ID INTEGER PRIMARY KEY AUTOINCREMENT,
            sWord TEXT,
            sType TEXT,
            isIgnore INTEGER,
            imgData BLOB
        );
        """
        self.cursor.execute(create_table_sql)
        self.conn.commit()
        #self.conn.close()
        pass
    
    def List_processed_documents(self):
        self.docx_files = [f for f in os.listdir(self.processed_data) if f.endswith('.docx')]
        
    def extract_image_from_run(self, run):
        drawing = run._element.xpath('.//w:drawing')
        if drawing:
            blips = run._element.xpath('.//a:blip')
            if blips:
                embed_rId = blips[0].get(qn('r:embed'))
                image_part = run.part.related_parts[embed_rId]
                return image_part.blob  # binary data
        return None    
        
    
    def OpenAll_PreProcess_Files(self):
        for aDocxFile in self.docx_files:
            sFileFull_Path = os.path.join(self.processed_data, aDocxFile)
            print(sFileFull_Path)
            #self.OpenDocx_ReadWords_by_Words(sFileFull_Path)
            self.OpenDocx_Read_Table_Data_From_Docx(sFileFull_Path)
            
            
    def Save_Image(self, index, sRun, image_blob):
        sFile = ''.join(random.sample(self.pool, 8))
        file_name = ""
        Image_Font = os.path.join(os.getcwd(), "Image_Font")
        if (os.path.exists(Image_Font) == False):
            os.makedirs(Image_Font)        
        if image_blob:
            file_name = os.path.join(Image_Font, f"{sFile}_{sRun}.jpg")
            with open(file_name, "wb") as f:
                f.write(image_blob)
        return file_name
    
    def convert_to_binary_data(self, filename):
        filename = filename.replace("\n", "")
        with open(filename, 'rb') as file:
            return file.read()    
                
    def Insert_Image_to_DB(self, sWord, imgData):
        SQL = f"INSERT INTO Word(sWord ,sType ,isIgnore, imgData) VALUES('{sWord}', 'A1', 0, ?)"
        isExists = self.dbCheck_Exist_B4_Insert("Word", sWord)
        if (isExists == None):        
            self.cursor.execute(SQL, (imgData, ))
            self.conn.commit()
            
    def Insert_Image_to_DB_Special_CASEA(self, sWord, imgData):
        SQL = f"INSERT INTO Word(sWord ,sType ,isIgnore, imgData) VALUES('{sWord}', 'A5', 0, ?)"
        isExists = self.dbCheck_Exist_B4_Insert("Word", sWord)
        if (isExists == None):        
            self.cursor.execute(SQL, (imgData, ))
            self.conn.commit()    
            
    def Insert_Image_to_DB_CASE_B(self, sWord, imgData):
        SQL = f"INSERT INTO Word(sWord ,sType ,isIgnore, imgData) VALUES('{sWord}', 'A1', 1, ?)"
        isExists = self.dbCheck_Exist_B4_Insert("Word", sWord)
        if (isExists == None):        
            self.cursor.execute(SQL, (imgData, ))
            self.conn.commit()    
        
    def dbCheck_Exist_B4_Insert(self, sTable, sWord):
        self.cursor.execute(f"SELECT * FROM {sTable} WHERE sWord = ?", (sWord,))
        return self.cursor.fetchone()
        
    
    def Insert_Sound_Words_to_DB(self, sWord):
        #normal_word
        SQL = f"INSERT INTO Word(sWord ,sType ,isIgnore, imgData) VALUES('{sWord}', 'A3', 0, NULL)"
        isExists = self.dbCheck_Exist_B4_Insert("Word", sWord)
        if (isExists == None):
            self.cursor.execute(SQL)
            self.conn.commit()
        
    def Insert_Normal_words_to_DB(self, sWord):
        SQL = f"INSERT INTO Word(sWord ,sType ,isIgnore, imgData) VALUES('{sWord}', 'normal', 0, NULL)"
        isExists = self.dbCheck_Exist_B4_Insert("Word", sWord)
        if (isExists == None):        
            self.cursor.execute(SQL)
            self.conn.commit()
            
    def CASE_A_Need_Highlight_Hard(self, index, para, run):
        #●	a1差異字（黃底＋截圖）編號1-編號43
        #●	a2差異字（待考究，黃底＋截圖）編號1-編號4

        ix = index + 1
        image_blob = self.extract_image_from_run(para.runs[ix])
        file_name = self.Save_Image(index, run.text, image_blob)
        binary_data = self.convert_to_binary_data(file_name)
        self.Insert_Image_to_DB(run.text, binary_data)        
        
    def CASE_Special_A_Need_Highlight_Hard_dual_sound(self, index, para, run):
        #a3差異字＋難字
        ix = index + 1
        image_blob = self.extract_image_from_run(para.runs[ix])
        file_name = self.Save_Image(index, run.text, image_blob)
        binary_data = self.convert_to_binary_data(file_name)
        self.Insert_Image_to_DB_Special_CASEA(run.text, binary_data)     
        
        
    def CASE_B_Need_Highlight_OK_Ignore(self, index, para, run):
        #【字庫B】有差異但是可以忽略的字→需刪除標示
        ix = index + 1
        image_blob = self.extract_image_from_run(para.runs[ix])
        file_name = self.Save_Image(index, run.text, image_blob)
        binary_data = self.convert_to_binary_data(file_name)
        self.Insert_Image_to_DB_CASE_B(run.text, binary_data)      

            
    def OpenDocx_Read_Table_Data_From_Docx(self, sFileFull_Path):
        document = Document(sFileFull_Path)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for index, run in enumerate(para.runs):
                            text = run.text.strip()
                            if (not text): continue
                            if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                                self.CASE_A_Need_Highlight_Hard(index, para, run)
                                print(f"Font diff {run.text}")
                            elif (run.font.highlight_color == WD_COLOR_INDEX.GRAY_25):
                                self.CASE_B_Need_Highlight_OK_Ignore(index, para, run)
                                print(f"Font diff, OK for Ignore {run.text}")
                            elif (run.font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN):
                                try:
                                    self.CASE_Special_A_Need_Highlight_Hard_dual_sound(index, para, run)
                                    print(f"Font + Dual {run.text}")
                                except:
                                    print(f"dual-sound--{run.text}")
                                    self.Insert_Sound_Words_to_DB(run.text)
                                
                                    
            
    def OpenDocx_ReadWords_by_Words(self, sFileFull_Path):
        docx = Document(sFileFull_Path)
        for para in docx.paragraphs:
            for index, run in enumerate(para.runs):
                #if (run.text in self.All_Processed_words):continue
                #self.All_Processed_words.append(run.text)
                if run.font.highlight_color:
                    print(f"{run.text}, {run.font.highlight_color}")
                    if (run.font.highlight_color == self.YELLOW):
                        ix = index + 1
                        image_blob = self.extract_image_from_run(para.runs[ix])
                        file_name = self.Save_Image(index, run.text, image_blob)
                        binary_data = self.convert_to_binary_data(file_name)
                        self.Insert_Image_to_DB(run.text, binary_data)
                        
                        
                    elif (run.font.highlight_color == self.BRIGHT_GREEN):
                        self.Insert_Sound_Words_to_DB(run.text)
                else:
                    for char in run.text:
                        char = char.strip()
                        self.Insert_Normal_words_to_DB(char)
        






def main():
    docx = docx1001()
    docx.List_processed_documents()
    docx.OpenAll_PreProcess_Files()
    pass


if __name__ =="__main__":
    main()
    
"""
將文字分成三類
自行
破音難字
一般文字
分別放入資料庫中
上面為第一步驟 建立資料這是第一隻程式

第二步驟搞成第二隻程式
讀取尚未處理文件
每個字拿出來讀取資料庫
如果在資料庫中已經有資料則標註相對應的type 難字或是字型
字型則需要符合規格加入圖片和背景色highlight
如果文字尚未出現在資料庫中，則必須要標示出來讓人類來檢查

第三支程式會有GUI
可以新增刪除修改查詢
已經建立的文字


"""