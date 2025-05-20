from docx import Document
from docx.oxml.ns import qn
import os
import sqlite3
import random
import string
from docx.enum.text import WD_COLOR_INDEX

from docx.shared import Inches
from docx.oxml import OxmlElement
import io


class Scan_Exists_Docx:
    def __init__(self):
        self.processed_data = os.path.join(".", "processed_document")
        self.docx_files = []
        self.conn = sqlite3.connect("word_data.db")
        self.cursor = self.conn.cursor()
        self.Create_Sqlite_DB()
        self.All_Processed_words = []
        self.YELLOW = 7
        self.BRIGHT_GREEN = 4
        self.pool = string.ascii_letters + string.digits
        self.A_Font_todo = None
        self.A_Dual_sound_todo = None
    
    def Create_Sqlite_DB(self):
        create_table_sql = """
        CREATE TABLE IF NOT EXISTS Word (
            ID INTEGER PRIMARY KEY AUTOINCREMENT,
            sWord TEXT,
            sType TEXT,
            isIgnore INTEGER,
            imgData BLOB
        );
        """
        self.cursor.execute(create_table_sql)
        self.conn.commit()
        #self.conn.close()
        pass
    
    def List_processed_documents(self):
        self.docx_files = [f for f in os.listdir(self.processed_data) if f.endswith('.docx')]
        
    def extract_image_from_run(self, run):
        drawing = run._element.xpath('.//w:drawing')
        if drawing:
            blips = run._element.xpath('.//a:blip')
            if blips:
                embed_rId = blips[0].get(qn('r:embed'))
                image_part = run.part.related_parts[embed_rId]
                return image_part.blob  # binary data
        return None    
        
    
    def OpenAll_PreProcess_Files(self):
        for aDocxFile in self.docx_files:
            sFileFull_Path = os.path.join(self.processed_data, aDocxFile)
            print(sFileFull_Path)
            self.OpenDocx_ReadWords_by_Words(sFileFull_Path)
            
            
    def Save_Image(self, index, sRun, image_blob):
        sFile = ''.join(random.sample(self.pool, 8))
        file_name = ""
        Image_Font = os.path.join(os.getcwd(), "Image_Font")
        if (os.path.exists(Image_Font) == False):
            os.makedirs(Image_Font)        
        if image_blob:
            file_name = os.path.join(Image_Font, f"{sFile}_{sRun}.jpg")
            with open(file_name, "wb") as f:
                f.write(image_blob)
        return file_name
    
    def convert_to_binary_data(self, filename):
        with open(filename, 'rb') as file:
            return file.read()    
                
    def Insert_Image_to_DB(self, sWord, imgData):
        SQL = f"INSERT INTO Word(sWord ,sType ,isIgnore, imgData) VALUES('{sWord}', 'FontType', 0, ?)"
        isExists = self.dbCheck_Exist_B4_Insert("Word", sWord)
        if (isExists == None):        
            self.cursor.execute(SQL, (imgData, ))
            self.conn.commit()
            
    def Insert_Image_to_DB_Special_CASEA(self, sWord, imgData):
        SQL = f"INSERT INTO Word(sWord ,sType ,isIgnore, imgData) VALUES('{sWord}', 'Font_Dual', 0, ?)"
        isExists = self.dbCheck_Exist_B4_Insert("Word", sWord)
        if (isExists == None):        
            self.cursor.execute(SQL, (imgData, ))
            self.conn.commit()    
            
    def Insert_Image_to_DB_CASE_B(self, sWord, imgData):
        SQL = f"INSERT INTO Word(sWord ,sType ,isIgnore, imgData) VALUES('{sWord}', 'FontType', 1, ?)"
        isExists = self.dbCheck_Exist_B4_Insert("Word", sWord)
        if (isExists == None):        
            self.cursor.execute(SQL, (imgData, ))
            self.conn.commit()    
        
    def dbCheck_Exist_B4_Insert(self, sTable, sWord):
        self.cursor.execute(f"SELECT * FROM {sTable} WHERE sWord = ?", (sWord,))
        return self.cursor.fetchone()
    def dbCheck_Font_ok_for_Ignore(self, sWord):
        SQL = f"Select * FROM [Word] WHERE isIgnore = 0 and sWord = '{sWord}'"
        self.cursor.execute(SQL)
        return self.cursor.fetchone()
    
    def dbCheck_Dual_Sound_Exists(self, sWord):
        SQL = f"Select * FROM [Word] WHERE sType = 'dual' and sWord = '{sWord}'"
        self.cursor.execute(SQL)
        return self.cursor.fetchone()        
    
    def Insert_Sound_Words_to_DB(self, sWord):
        #normal_word
        SQL = f"INSERT INTO Word(sWord ,sType ,isIgnore, imgData) VALUES('{sWord}', 'dual', 0, NULL)"
        isExists = self.dbCheck_Exist_B4_Insert("Word", sWord)
        if (isExists == None):
            self.cursor.execute(SQL)
            self.conn.commit()
        
    def Insert_Normal_words_to_DB(self, sWord):
        SQL = f"INSERT INTO Word(sWord ,sType ,isIgnore, imgData) VALUES('{sWord}', 'normal', 0, NULL)"
        isExists = self.dbCheck_Exist_B4_Insert("Word", sWord)
        if (isExists == None):        
            self.cursor.execute(SQL)
            self.conn.commit()
            
    def CASE_A_Need_Highlight_Hard(self, index, para, run):
        #●	a1差異字（黃底＋截圖）編號1-編號43
        #●	a2差異字（待考究，黃底＋截圖）編號1-編號4

        ix = index + 1
        image_blob = self.extract_image_from_run(para.runs[ix])
        file_name = self.Save_Image(index, run.text, image_blob)
        binary_data = self.convert_to_binary_data(file_name)
        self.Insert_Image_to_DB(run.text, binary_data)        
        
    def CASE_Special_A_Need_Highlight_Hard_dual_sound(self, index, para, run):
        #a3差異字＋難字
        ix = index + 1
        image_blob = self.extract_image_from_run(para.runs[ix])
        file_name = self.Save_Image(index, run.text, image_blob)
        binary_data = self.convert_to_binary_data(file_name)
        self.Insert_Image_to_DB_Special_CASEA(run.text, binary_data)     
        
        
    def CASE_B_Need_Highlight_OK_Ignore(self, index, para, run):
        #【字庫B】有差異但是可以忽略的字→需刪除標示
        ix = index + 1
        image_blob = self.extract_image_from_run(para.runs[ix])
        file_name = self.Save_Image(index, run.text, image_blob)
        binary_data = self.convert_to_binary_data(file_name)
        self.Insert_Image_to_DB_CASE_B(run.text, binary_data)      

            
    def OpenDocx_Read_Table_Data_From_Docx(self, sFileFull_Path):
        document = Document(sFileFull_Path)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for index, run in enumerate(para.runs):
                            text = run.text.strip()
                            if (not text): continue
                            if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                                self.CASE_A_Need_Highlight_Hard(index, para, run)
                                print(f"Font diff {run.text}")
                            elif (run.font.highlight_color == WD_COLOR_INDEX.TURQUOISE):
                                self.CASE_B_Need_Highlight_OK_Ignore(index, para, run)
                                print(f"Font diff, OK for Ignore {run.text}")
                            elif (run.font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN):
                                try:
                                    self.CASE_Special_A_Need_Highlight_Hard_dual_sound(index, para, run)
                                    print(f"Font + Dual {run.text}")
                                except:
                                    print(f"dual-sound--{run.text}")
                                    self.Insert_Sound_Words_to_DB(run.text)
                                
                                    
    def Create_Docx_if_not_Exists(self, sDocxFile_Path):
        if (os.path.exists(sDocxFile_Path) == True):
            if ("A_Font_todo" in sDocxFile_Path):
                self.A_Font_todo = Document(sDocxFile_Path)
            elif ("A_Dual_sound_todo" in sDocxFile_Path):
                self.A_Dual_sound_todo = Document(sDocxFile_Path)
        else:
            if (os.path.exists(sDocxFile_Path) == False):
                if ("A_Font_todo" in sDocxFile_Path):
                    self.A_Font_todo = Document()
                elif ("A_Dual_sound_todo" in sDocxFile_Path):
                    self.A_Dual_sound_todo = Document()


                
        
            
    def Create_Todo_Files(self):
        self.A_Font_todo_File = "A_Font_todo.docx"
        self.A_Dual_sound_todo_File = "A_Dual_sound_todo.docx"
        lstFiles = [self.A_Font_todo_File, self.A_Dual_sound_todo_File]
        for aDocx in lstFiles:
            self.Create_Docx_if_not_Exists(aDocx)
    
    def highlight_run(self, run, color):
        #new_run = para.add_run(ch)
        run.font.highlight_color = color  # TURQUOISE
        #new_runs.append(new_run)        
        """
        highlight = OxmlElement('w:highlight')
        highlight.set(qn('w:val'), color)
        run._r.get_or_add_rPr().append(highlight)
        """
        
    def get_highlight_color(self, run):
        """回傳 run 的 highlight 顏色字串（如 'yellow'）"""
        hl = run._element.xpath(".//w:highlight")
        if hl:
            return hl[0].get(qn("w:val"))
        return None
    
    def set_highlight_color(self, run, color):
        """將 run 的 highlight 設為指定顏色"""
        highlight = OxmlElement("w:highlight")
        highlight.set(qn("w:val"), color)
        rPr = run._r.get_or_add_rPr()
        # 移除舊 highlight
        for el in rPr.findall(qn("w:highlight")):
            rPr.remove(el)
        rPr.append(highlight)    
        
    def insert_highlight_and_image(self, doc, char, image_bytes, color=WD_COLOR_INDEX.YELLOW):
        para = doc.add_paragraph()
        run = para.add_run(char)
        self.highlight_run(run, color=WD_COLOR_INDEX.YELLOW)
        image_stream = io.BytesIO(image_bytes)
        #para.add_run().add_picture(image_stream, width=Inches(1.5))
        para.add_run().add_picture(image_stream)
        
    def search_in_docx(self, doc, keyword):
        results = []
        for para in doc.paragraphs:
            if keyword in para.text:
                results.append(para.text)
        return results    
        
            
    def Insert_Record_A_Font_todo(self, sWord, bImg):
        res = self.search_in_docx(self.A_Font_todo, sWord)
        if (res == []):
            self.insert_highlight_and_image(self.A_Font_todo, sWord, bImg)
        
    def Insert_Record_A_Dual_todo(self, sWord):
        res = self.search_in_docx(self.A_Dual_sound_todo, sWord)
        if (res == []):        
            para = self.A_Dual_sound_todo.add_paragraph()
            run = para.add_run(sWord)
            self.highlight_run(run, color=WD_COLOR_INDEX.BRIGHT_GREEN)
    
    def Word_IN_CASE_B_DB(self, sWord):
        pass
        
    
    def OpenDocx_ReadWords_by_Words(self, sFileFull_Path):
        docx = Document(sFileFull_Path)
        self.Create_Todo_Files()
        for para in docx.paragraphs:
            for index, run in enumerate(para.runs):
                hl_color = self.get_highlight_color(run)
                if hl_color == "yellow":
                    Ignore_word = self.dbCheck_Font_ok_for_Ignore(run.text)
                    if (Ignore_word != None):
                        self.highlight_run(run, WD_COLOR_INDEX.TURQUOISE)
                    else:
                        ix = index + 1
                        image_blob = self.extract_image_from_run(para.runs[ix])
                        file_name = self.Save_Image(index, run.text, image_blob)
                        binary_data = self.convert_to_binary_data(file_name)                        
                        self.Insert_Record_A_Font_todo(run.text, binary_data)
                elif (hl_color == "green"):
                    dual_sound = self.dbCheck_Dual_Sound_Exists(run.text)
                    if (dual_sound != None): continue
                    self.Insert_Record_A_Dual_todo(run.text)
                    
                        
                        
                    
        docx.save(sFileFull_Path)
               
        self.A_Font_todo.save(self.A_Font_todo_File)
        self.A_Dual_sound_todo.save(self.A_Dual_sound_todo_File)
        print(os.path.join(os.getcwd(), self.A_Font_todo_File))
        print(os.path.join(os.getcwd(), self.A_Dual_sound_todo_File))







def main():
    docx = Scan_Exists_Docx()
    docx.List_processed_documents()
    docx.OpenAll_PreProcess_Files()
    pass


if __name__ =="__main__":
    main()
    
"""
Open Exists docx file.
compare Yellow Highlight and dual sound with Database.
if the word not exists in DB
export to another docx file with file name
A_Font_todo.docx => Yellow Highligh with image
A_Dual_sound_todo.docx => dual sound.
"""